import fitz  # PyMuPDF
from docx import Document
import re
import logging
from pathlib import Path
from typing import Union

logger = logging.getLogger(__name__)


def extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return clean_text(text)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        import io
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return clean_text("\n".join(paragraphs))
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text(file_bytes: bytes) -> str:
    """Extract from plain text."""
    try:
        return clean_text(file_bytes.decode("utf-8", errors="ignore"))
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return ""


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove special chars but keep useful punctuation
    text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\[\]\@\#\/\+]', ' ', text)
    return text.strip()


def ingest_file(file_bytes: bytes, filename: str) -> str:
    """Route file to correct extractor based on extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_docx(file_bytes)
    elif ext in (".txt", ".md"):
        return extract_text(file_bytes)
    else:
        # Try as text fallback
        return extract_text(file_bytes)